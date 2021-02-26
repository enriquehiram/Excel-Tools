import openpyxl
import arrow
from docx import Document
from docx.shared import Inches

reng_ini = 2
reng_fin = 14
col_ini = 0
col_fin = 9
col = ['B','C','D','E','F','G','H','I','J']

wb = openpyxl.load_workbook('Minuta.xlsx')
ws = wb["Minutable"]

document = Document()
document.add_heading('Minute', 0)
p = document.add_paragraph('Bellow the meeting minute updated. Please answer against this e-mail with your feedback if doubts, corrections or adds if any.')
q = document.add_paragraph ('If during the pending point is “INFO” stated, this means this topic is just informative, therefore, no additional task or answer is required from you.')


for j in range (col_ini, col_fin,1):
	team = ws[ col[j] + "1"].value
	document.add_heading( team, level = 1)
	for i in range (reng_ini, reng_fin + 1, 1):
		program = ws["A" + str(i)].value
		current_cel = ws[ col[j] + str(i) ].value
		if current_cel != None :
			actividades_cel = current_cel.split('\n\n')
			for k in range(len(actividades_cel)):
				document.add_paragraph(program + " " + actividades_cel[k], style='List Bullet')


# for i in range (reng_ini, reng_fin + 1, 1):
	# program = ws["A" + str(i)].value
	# current_cel = ws[ col[0] + str(i) ].value
	# if current_cel != None :
		# actividades_cel = current_cel.split('\n\n')
		# for k in range(len(actividades_cel)):
			# document.add_paragraph(program + " " + actividades_cel[k], style='List Bullet')

document.save('minuta.docx')
